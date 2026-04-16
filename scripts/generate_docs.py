"""Generate/update AgentSquad executive documents."""
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

DOCS_DIR = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), "docs")


def add_title(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor(0, 51, 102)


def add_subtitle(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(102, 102, 102)


def add_h1(doc, text):
    p = doc.add_heading(text, level=1)
    for run in p.runs:
        run.font.color.rgb = RGBColor(0, 51, 102)


def add_h2(doc, text):
    p = doc.add_heading(text, level=2)
    for run in p.runs:
        run.font.color.rgb = RGBColor(0, 80, 140)


def add_h3(doc, text):
    p = doc.add_heading(text, level=3)
    for run in p.runs:
        run.font.color.rgb = RGBColor(0, 100, 160)


def add_para(doc, text):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(8)
    return p


def add_bullet(doc, text):
    return doc.add_paragraph(text, style="List Paragraph")


def add_bold_para(doc, label, text):
    p = doc.add_paragraph()
    run = p.add_run(label)
    run.bold = True
    p.add_run(text)
    p.paragraph_format.space_after = Pt(8)


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            table.rows[ri + 1].cells[ci].text = val
    doc.add_paragraph()


# ============================================================
# DOCUMENT 1: Update Executive Summary
# ============================================================
def generate_executive_summary():
    doc = Document()

    add_title(doc, "AgentSquad: Human-Agent Hybrid Development")
    add_subtitle(doc, "A Proposal for AI-Augmented Software Engineering at Microsoft")
    add_subtitle(doc, "Prepared by: Ben Humphrey  |  Date: April 2026  |  CONFIDENTIAL")
    doc.add_paragraph()

    add_h1(doc, "Executive Summary")
    add_para(doc,
        "Software development is undergoing a fundamental transformation. AI coding assistants "
        "like GitHub Copilot have proven that AI can boost individual developer productivity. But "
        "the real opportunity lies at the team level \u2014 what if an entire development team could "
        "be AI-augmented, with human engineers directing AI agents through the full software "
        "development lifecycle?")
    add_para(doc,
        "AgentSquad is a working, battle-tested prototype that demonstrates this future. It is a "
        "multi-agent AI system where 7+ specialized AI agents \u2014 Program Manager, Researcher, "
        "Architect, Software Engineer, Software Engineers, and Test Engineer \u2014 "
        "collaborate through real GitHub PRs and Issues to build software end-to-end. Over 85+ "
        "iterative build-run-fix cycles, the system has been hardened with crash recovery, rate "
        "limiting, vision-based PR review, multi-tier test automation (including Playwright UI "
        "tests), and 23 documented operational lessons.")
    add_para(doc,
        "Since the initial prototype, AgentSquad has evolved significantly with capabilities "
        "that dramatically increase its power and flexibility:")
    add_bullet(doc,
        "Dynamic SME Agents \u2014 The PM and SE can spawn Subject Matter Expert agents on-demand "
        "(security auditors, database specialists, etc.) with custom AI personas, MCP tool "
        "servers, and external knowledge sources")
    add_bullet(doc,
        "Prompt Externalization \u2014 All ~95 agent prompts live in editable .md template files, "
        "allowing behavior tuning without code changes or recompilation")
    add_bullet(doc,
        "MCP Server Integration \u2014 Agents can be equipped with Model Context Protocol tool "
        "servers for code search, documentation, and external API access")
    add_bullet(doc,
        "Knowledge Pipeline \u2014 Agents automatically fetch, extract, and summarize external "
        "documentation URLs, injecting domain expertise into their prompts")
    add_bullet(doc,
        "Custom Agent Definitions \u2014 New agent roles defined via configuration alone (persona, "
        "tools, knowledge links) without writing code")
    add_bullet(doc,
        "15-Page Real-Time Dashboard \u2014 Blazor Server UI with agent overview, project timeline, "
        "metrics, health monitor, engineering plan dependency graph, team visualization, director "
        "CLI terminal, agent reasoning chains, and approval management")
    add_bullet(doc,
        "Incremental Code Generation \u2014 Agents now read existing file content before modifying, "
        "making surgical changes rather than full rewrites")
    add_para(doc,
        "The proposed evolution transforms AgentSquad from a fully autonomous pipeline into a "
        "human-agent hybrid development platform where team members inject themselves into the "
        "virtual team at any level \u2014 acting as director of the whole team, or pairing with "
        "their role-matched agent.")

    add_h1(doc, "The Opportunity")
    add_para(doc,
        "Microsoft engineering teams face persistent challenges: growing backlogs, staffing "
        "constraints, and pressure to ship faster while maintaining quality. Current AI tools "
        "help individual developers write code faster, but they don't address the systemic "
        "challenges of coordinating complex software projects across multiple disciplines.")
    add_para(doc, "AgentSquad addresses this at the team level:")
    add_bullet(doc,
        "A single developer acting as \"Director\" can manage an entire virtual dev team of 7+ "
        "AI agents, potentially leveraging AI for 85% of development and guiding the remaining "
        "15% with human judgment")
    add_bullet(doc,
        "Alternatively, each team member pairs with their role-matched agent (PM human + PM agent, "
        "Architect human + Architect agent, etc.) \u2014 the agent handles execution while the "
        "human provides strategic direction")
    add_bullet(doc,
        "Dynamic SME agents extend the team on-demand: need a security audit? The PM spawns a "
        "SecurityAuditor SME with OWASP knowledge. Need database expertise? Spawn a "
        "DatabaseArchitect SME with schema docs")
    add_bullet(doc,
        "Agents handle 85%+ of execution: research, specs, architecture, coding, testing, reviews")
    add_bullet(doc,
        "Humans focus on what they do best: decisions, judgment calls, creative direction, "
        "stakeholder alignment")
    add_bullet(doc,
        "All work products live in GitHub (PRs, Issues, reviews) \u2014 fully auditable, no vendor lock-in")
    add_bullet(doc,
        "Externalized prompts allow rapid iteration on agent behavior without engineering cycles")

    add_h1(doc, "How It Works")

    add_h2(doc, "Graduated Autonomy")
    add_para(doc,
        "The key innovation is configurable human oversight at every decision point. Three "
        "presets make this simple:")
    add_bullet(doc, "Full Auto \u2014 Agents run autonomously end-to-end (good for low-risk features)")
    add_bullet(doc,
        "Supervised \u2014 Human approval at critical gates (spec, architecture, final PR merge)")
    add_bullet(doc,
        "Full Control \u2014 Human approval at every gate (maximum oversight for critical systems)")
    add_para(doc,
        "Gate configuration is hot-reloadable \u2014 change oversight level at runtime without "
        "restarting. The human \"Director\" can approve, reject with feedback (triggering automatic "
        "rework), or modify proposals through the dashboard.")

    add_h2(doc, "AI-Powered Team Composition")
    add_para(doc,
        "The PM agent doesn't just manage tasks \u2014 it designs the team. After reading the "
        "project description and research, the PM proposes an optimal team: how many engineers, "
        "which specialists (SME agents) are needed, what MCP tool servers each agent should have, "
        "and what external knowledge to ingest. The human director reviews and approves the team "
        "composition before agents are spawned.")

    add_h2(doc, "GitHub Copilot CLI as AI Backend")
    add_para(doc,
        "All AI interactions route through the GitHub Copilot CLI binary by default \u2014 no "
        "separate API keys required. This leverages existing Copilot licenses with automatic "
        "fallback to direct API providers. The system supports Anthropic Claude, OpenAI GPT, "
        "Azure OpenAI, and local Ollama models across four quality tiers (premium, standard, "
        "budget, local).")

    add_h2(doc, "15-Page Real-Time Dashboard")
    add_para(doc,
        "A Blazor Server dashboard provides complete visibility: agent overview with status "
        "badges, project timeline with phase grouping, system metrics, health monitoring with "
        "stuck-agent detection, PR and Issue browsers, an interactive engineering plan dependency "
        "graph, team visualization, a Director CLI terminal, approval management, prompt template "
        "editor, agent reasoning chain viewer, and GitHub activity feed.")

    add_h1(doc, "Benefits for Microsoft Engineering")
    add_bold_para(doc, "Productivity Multiplication: ",
        "A single developer directing AgentSquad can produce the output equivalent of a small "
        "team. Agents handle research, spec writing, architecture, coding, testing, and reviews "
        "autonomously \u2014 the human provides direction and quality gates.")
    add_bold_para(doc, "Quality Through Systematic Review: ",
        "Every PR gets multi-agent peer review (PM for requirements, SE for architecture, TE for "
        "test coverage). Vision-based screenshot review catches UI regressions. Incremental "
        "modification rules prevent agents from rewriting existing code during feature additions.")
    add_bold_para(doc, "Extensible Expertise: ",
        "Dynamic SME agents with MCP tool servers and knowledge pipelines mean the system's "
        "capabilities grow with each project. A SecurityAuditor SME created for one project "
        "becomes a reusable template for future projects.")
    add_bold_para(doc, "Knowledge Continuity: ",
        "Agents don't forget context, go on vacation, or need onboarding. SQLite-backed memory "
        "persists decisions and learnings across restarts. 23 documented lessons from 85+ runs "
        "create institutional knowledge.")
    add_bold_para(doc, "Scalable Oversight: ",
        "The graduated autonomy model means humans aren't bottlenecks. Low-risk work flows "
        "autonomously; high-risk decisions pause for review. Hot-reloadable gate configuration "
        "adapts oversight in real time.")
    add_bold_para(doc, "GitHub-Native: ",
        "All artifacts live in GitHub repositories \u2014 no proprietary platform required. Teams "
        "can adopt AgentSquad alongside their existing workflows.")
    add_bold_para(doc, "Prompt-Driven Iteration: ",
        "~95 externalized prompt templates in editable .md files allow rapid agent behavior "
        "tuning without code changes, recompilation, or redeployment.")

    add_h1(doc, "Implementation Status")
    add_para(doc, "AgentSquad is not a proposal \u2014 it is a working system with:")
    add_bullet(doc, "7 core agent roles + dynamic SME agents + custom agent definitions")
    add_bullet(doc, "167+ passing automated tests (unit, integration, agent behavior)")
    add_bullet(doc, "15-page Blazor Server dashboard with real-time SignalR updates")
    add_bullet(doc, "~95 externalized prompt templates with hot-reload and dashboard editor")
    add_bullet(doc, "MCP server integration, knowledge pipeline, AI-powered team composition")
    add_bullet(doc, "Multi-tier test automation (unit, integration, Playwright UI/E2E)")
    add_bullet(doc, "Crash-resilient sessions with SQLite state persistence")
    add_bullet(doc, "GitHub API rate limiting with proactive throttling (~90% call reduction)")
    add_bullet(doc, "23 documented operational lessons from 85+ iterative runs")
    add_bullet(doc, "Comprehensive requirements document (30 sections, 1800+ lines)")

    add_h1(doc, "The Vision: Human-Directed AI Development Teams")
    add_para(doc,
        "The ultimate vision for AgentSquad is a development model where human engineers work "
        "alongside AI agents as naturally as they work alongside human colleagues:")
    add_bullet(doc,
        "Solo Director Mode \u2014 One developer manages the entire agent team, reviewing specs, "
        "approving architectures, and guiding implementation. The developer provides 15% direction; "
        "AI handles 85% execution")
    add_bullet(doc,
        "Role Pairing Mode \u2014 Each human team member pairs with their agent counterpart. The "
        "human PM directs the PM agent, the human architect reviews the Architect agent's designs, "
        "engineers guide their engineer agents. Each person amplifies their output through their "
        "AI partner")
    add_bullet(doc,
        "Hybrid Team Mode \u2014 A mix of both \u2014 some roles filled by human-agent pairs, others "
        "fully automated. A small 3-person team could operate with the throughput of a 15-person team")
    add_para(doc,
        "In all modes, the key insight is the same: AI agents handle the volume work (research, "
        "first drafts, coding, test writing, review mechanics) while humans provide the "
        "irreplaceable elements (strategic decisions, creative direction, judgment calls, "
        "stakeholder communication). The result is not just faster development \u2014 it is a "
        "fundamentally different development model where the human's role evolves from maker to director.")

    add_h1(doc, "Recommendation")
    add_para(doc,
        "We recommend a funded proof-of-concept applying the human-agent hybrid model to a real "
        "internal project. This would:")
    add_bullet(doc, "Validate the productivity multiplier claims with measurable data")
    add_bullet(doc, "Identify friction points in the human-agent collaboration model")
    add_bullet(doc, "Produce a reusable platform that other Microsoft teams can adopt")
    add_bullet(doc,
        "Position Microsoft's internal engineering as a showcase for AI-augmented development")
    add_para(doc,
        "The AI-augmented development model is not theoretical \u2014 AgentSquad is already building "
        "software autonomously with human oversight. The next step is to bring a human team into "
        "the loop and measure the impact.")
    doc.add_paragraph()
    add_para(doc, "For questions or discussion, contact Ben Humphrey (behumphr@microsoft.com)")

    path = os.path.join(DOCS_DIR, "AgentSquad-Executive-Summary.docx")
    doc.save(path)
    print(f"Executive Summary saved to {path}")


# ============================================================
# DOCUMENT 2: Detailed Project Summary
# ============================================================
def generate_detailed_summary():
    doc = Document()

    add_title(doc, "AgentSquad: Complete Solution Guide")
    add_subtitle(doc, "Architecture, Features, Capabilities, and Differentiation")
    add_subtitle(doc, "Prepared by: Ben Humphrey  |  Date: April 2026  |  CONFIDENTIAL")
    doc.add_paragraph()

    # --- TOC ---
    add_h1(doc, "Table of Contents")
    toc_items = [
        "1. Solution Overview",
        "2. Architecture & Design",
        "3. Agent Roles & Responsibilities",
        "4. Complete Feature Inventory",
        "5. The Development Workflow",
        "6. Human-Agent Collaboration Model",
        "7. Dashboard & Monitoring",
        "8. AI Provider Strategy",
        "9. Quality & Safety Engineering",
        "10. How to Use AgentSquad",
        "11. Benefits & Value Proposition",
        "12. Competitive Differentiation",
        "13. Vision & Roadmap",
    ]
    for item in toc_items:
        add_bullet(doc, item)

    # --- Section 1 ---
    add_h1(doc, "1. Solution Overview")
    add_para(doc,
        "AgentSquad is a .NET 8 multi-agent AI system that manages a full software development "
        "team \u2014 from Program Manager through Test Engineer \u2014 to autonomously build software "
        "projects. You provide a project description and a GitHub repository; AgentSquad handles "
        "research, architecture, engineering planning, parallel implementation, multi-tier testing, "
        "code review, and delivery.")
    add_para(doc,
        "Every artifact lives in GitHub as real PRs and Issues. A Blazor Server dashboard gives "
        "real-time visibility, and configurable human gates let you control how much autonomy the "
        "team has \u2014 from fully autonomous to fully supervised.")

    add_h2(doc, "What Makes It Different")
    add_para(doc,
        "Unlike single-agent coding tools (GitHub Copilot Workspace, Cursor, Cline) that help "
        "one developer write code faster, AgentSquad simulates an entire development team with "
        "role specialization, peer review, and workflow coordination. It is not a code generator "
        "\u2014 it is a team simulator that produces auditable, reviewed, tested software through "
        "real GitHub workflows.")

    add_h2(doc, "Key Numbers")
    add_table(doc,
        ["Metric", "Value"],
        [
            ["Core agent roles", "5 (PM, Researcher, Architect, SE, TE)"],
            ["Dynamic agent types", "SME Agents + Custom Agents (unlimited definitions)"],
            ["Automated tests", "167+ (unit, integration, agent behavior)"],
            ["Dashboard pages", "15 (real-time Blazor Server + SignalR)"],
            ["Prompt templates", "~95 externalized .md files"],
            ["Requirements documented", "30 sections, 1800+ lines"],
            ["Operational lessons", "23 documented from 85+ runs"],
            ["Lines of C# code", "~40,000+ across 5 projects"],
            ["Known bugs fixed", "25+ (documented in appendix)"],
        ])

    # --- Section 2 ---
    add_h1(doc, "2. Architecture & Design")

    add_h2(doc, "Project Structure")
    add_para(doc,
        "AgentSquad is organized as a multi-project .NET solution with clear separation of concerns:")
    add_table(doc,
        ["Project", "Purpose"],
        [
            ["AgentSquad.Core", "Shared abstractions: IAgent, message bus, GitHub service, "
             "persistence, Semantic Kernel integration, Copilot CLI provider, MCP config, "
             "knowledge pipeline"],
            ["AgentSquad.Agents", "7 concrete agent implementations + CustomAgent + SmeAgent + AgentFactory"],
            ["AgentSquad.Orchestrator", "Runtime coordination: workflow state machine, agent registry, "
             "spawn manager, deadlock detector, health monitor, graceful shutdown"],
            ["AgentSquad.Runner", "Application host: DI registration, REST API, bootstrap worker"],
            ["AgentSquad.Dashboard", "Blazor Server UI: 15 pages, SignalR hub, data services"],
            ["AgentSquad.Dashboard.Host", "Standalone dashboard process (port 5051)"],
        ])

    add_h2(doc, "Dual-Layer Communication")
    add_para(doc,
        "Agents communicate through two complementary layers simultaneously:")
    add_bold_para(doc, "In-Process Message Bus: ",
        "System.Threading.Channels with bounded capacity (1000 messages). Sub-millisecond "
        "delivery for real-time agent-to-agent coordination. Seven message types: TaskAssignment, "
        "StatusUpdate, HelpRequest, ResourceRequest, ReviewRequest, SpawnSme, SmeResult. "
        "Messages route by agent ID \u2014 set to \"*\" for broadcast.")
    add_bold_para(doc, "GitHub API: ",
        "Durable artifacts via Octokit.net. PRs for code, Issues for tasks, Comments for "
        "discussion, Labels for workflow state. 60-second TTL cache reduces API calls by ~90%. "
        "Rate limiting with proactive throttling (slowdown at <200 remaining, pause at <10).")

    add_h2(doc, "Phase-Gated Workflow")
    add_para(doc,
        "A WorkflowStateMachine enforces linear phase progression with no backward transitions:")
    add_para(doc,
        "Initialization \u2192 Research \u2192 Architecture \u2192 Engineering Planning \u2192 "
        "Parallel Development \u2192 Testing \u2192 Review \u2192 Finalization")
    add_para(doc,
        "Each phase has gate conditions (signals + document readiness) that must be met before "
        "advancing. Signals use dot-notation naming (e.g., \"research.complete\", "
        "\"engineering.plan.ready\"). Human gates can be inserted at any phase transition.")

    add_h2(doc, "Document Flow")
    add_para(doc,
        "The agent pipeline produces shared Markdown documents that each phase builds upon:")
    add_para(doc,
        "Project Description \u2192 Research.md (Researcher) \u2192 PMSpec.md (PM) \u2192 "
        "Architecture.md (Architect, reviewed by SE) \u2192 EngineeringPlan.md (SE) \u2192 "
        "TeamComposition.md (PM) \u2192 PRs with code (Engineers) \u2192 Test PRs (TE)")

    add_h2(doc, "Thread Safety")
    add_para(doc,
        "AgentRegistry uses ConcurrentDictionary for lock-free reads. WorkflowStateMachine and "
        "AgentSpawnManager use lock for state transitions. DeadlockDetector snapshots its graph "
        "before DFS traversal. AgentBase.Status has a dedicated status lock.")

    # --- Section 3 ---
    add_h1(doc, "3. Agent Roles & Responsibilities")

    add_h2(doc, "Core Team (Always Present)")
    add_table(doc,
        ["Role", "Model Tier", "Key Responsibilities"],
        [
            ["Program Manager", "Premium", "Orchestrates team composition, writes PMSpec with user stories, "
             "triages blockers, reviews PRs for business alignment, manages escalations to human executive, "
             "proposes and spawns SME agents"],
            ["Researcher", "Standard", "Multi-turn technical research, technology evaluation, "
             "feasibility analysis. Produces Research.md"],
            ["Architect", "Premium", "System design via 5-turn AI conversation, API/data modeling, "
             "technology selection. Produces Architecture.md. Reviews PRs for architectural compliance"],
            ["Software Engineer", "Premium", "Decomposes architecture into engineering tasks, assigns "
             "work by complexity, conducts rigorous code reviews with scoring rubrics, handles "
             "high-complexity PRs, creates integration PR, spawns reactive SME agents"],
            ["Software Engineer", "Standard", "Implements medium-complexity tasks with plan \u2192 implement "
             "\u2192 self-review pipeline. Local build/test verification before PR submission"],
            ["Software Engineer", "Budget", "Implements low-complexity tasks with self-validation retries. "
             "Escalates tasks that exceed capability threshold"],
            ["Test Engineer", "Standard", "Three-tier test generation (unit \u2192 integration \u2192 UI/E2E), "
             "testability assessment, source-bug classification, Playwright screenshots, coverage tracking"],
        ])

    add_h2(doc, "Dynamic Specialists")
    add_table(doc,
        ["Type", "Creation", "Lifecycle"],
        [
            ["Custom Agents", "Defined in config (persona, MCP servers, knowledge links)", "Persistent"],
            ["SME Agents", "AI-generated or from templates when specialist knowledge needed",
             "OnDemand, Continuous, or OneShot"],
            ["Additional Engineers", "PM requests scaling; Orchestrator enforces limits", "Persistent"],
        ])

    add_h2(doc, "SME Agent System")
    add_para(doc,
        "The SME (Subject Matter Expert) system is one of AgentSquad's most innovative features. "
        "When the PM analyzes a project, it determines what specialist expertise is needed beyond "
        "the core team. It can spawn SME agents from pre-configured templates or generate entirely "
        "new agent definitions using AI \u2014 complete with custom personas, MCP tool servers for "
        "specialized capabilities, and external knowledge sources for domain expertise.")
    add_para(doc, "Safety controls include:")
    add_bullet(doc, "Maximum 5 concurrent SME agents (configurable)")
    add_bullet(doc, "Per-definition instance limits")
    add_bullet(doc, "Human gate approval before spawning")
    add_bullet(doc, "Only PM and SE can spawn SMEs (no agent-spawns-agent chains)")
    add_bullet(doc, "Metrics tracking for all spawn/retirement events")

    # --- Section 4 ---
    add_h1(doc, "4. Complete Feature Inventory")

    add_h2(doc, "Core Development Features")
    add_bullet(doc, "Full SDLC Automation \u2014 Research through delivery in a single pipeline")
    add_bullet(doc, "Multi-Turn AI Conversations \u2014 Agents use Semantic Kernel ChatHistory for "
                     "stateful multi-turn conversations (Researcher: 3 turns, Architect: 5 turns)")
    add_bullet(doc, "Incremental Code Generation \u2014 Agents read existing files before modifying, "
                     "making surgical additions rather than full rewrites")
    add_bullet(doc, "Local Build & Test Verification \u2014 Real dotnet build/test in isolated workspaces "
                     "before any code reaches GitHub")
    add_bullet(doc, "Multi-Step Implementation \u2014 Engineers break tasks into 3-6 discrete steps, "
                     "each committed atomically")
    add_bullet(doc, "Self-Review Pipeline \u2014 Software Engineers review their own code before submission")

    add_h2(doc, "Testing Features")
    add_bullet(doc, "Three-Tier Test Generation \u2014 Unit \u2192 Integration \u2192 UI/E2E (Playwright)")
    add_bullet(doc, "AI Failure Classification \u2014 Distinguishes test bugs from source bugs")
    add_bullet(doc, "Automatic Fix-Retry Loops \u2014 AI reads failures, generates fixes, retests")
    add_bullet(doc, "Playwright UI Testing \u2014 Browser install, app lifecycle, screenshot capture")
    add_bullet(doc, "Port Isolation \u2014 Each agent gets a unique port via workspace path hashing")
    add_bullet(doc, "Vision-Based PR Review \u2014 Downloads and analyzes screenshots as base64 images")

    add_h2(doc, "Team Management Features")
    add_bullet(doc, "AI-Powered Team Composition \u2014 PM proposes optimal team based on project analysis")
    add_bullet(doc, "Dynamic SME Agents \u2014 On-demand specialists with custom personas and tools")
    add_bullet(doc, "Custom Agent Definitions \u2014 New roles via configuration, no code required")
    add_bullet(doc, "Resource Scaling \u2014 PM requests additional engineers; Orchestrator enforces limits")
    add_bullet(doc, "Agent Memory \u2014 SQLite-backed persistent recall across restarts (30 recent entries)")

    add_h2(doc, "Review & Quality Features")
    add_bullet(doc, "Multi-Agent Peer Review \u2014 PM (requirements) + SE (architecture) + TE (tests)")
    add_bullet(doc, "Configurable Rework Cycles \u2014 Max review rounds with force-approval fallback")
    add_bullet(doc, "Integration PR \u2014 SE creates final assembly PR after all tasks merge")
    add_bullet(doc, "Conflict Resolution \u2014 Detects and resolves merge conflicts across agent PRs")
    add_bullet(doc, "Incremental Modification Check \u2014 SE review flags excessive file rewrites")

    add_h2(doc, "Infrastructure Features")
    add_bullet(doc, "Prompt Externalization \u2014 ~95 templates in .md files with {{variable}} substitution")
    add_bullet(doc, "MCP Server Integration \u2014 Code search, documentation, external APIs via tool servers")
    add_bullet(doc, "Knowledge Pipeline \u2014 Fetch, extract, summarize external docs into agent prompts")
    add_bullet(doc, "GitHub API Rate Limiting \u2014 Proactive throttling, smart reset-timestamp pausing")
    add_bullet(doc, "60s TTL API Cache \u2014 ~90% reduction in GitHub API calls")
    add_bullet(doc, "Crash-Resilient Sessions \u2014 CLI session IDs persist across restarts")
    add_bullet(doc, "Deadlock Detection \u2014 DFS wait-for graph analysis with cycle detection")
    add_bullet(doc, "Health Monitoring \u2014 Stuck agent detection, system diagnostics")
    add_bullet(doc, "Graceful Shutdown \u2014 State persistence, clean agent lifecycle management")

    add_h2(doc, "Human Oversight Features")
    add_bullet(doc, "Configurable Human Gates \u2014 17+ natural approval points in the workflow")
    add_bullet(doc, "Three Presets \u2014 Full Auto, Supervised, Full Control")
    add_bullet(doc, "Hot-Reloadable Config \u2014 Change gate settings at runtime, no restart needed")
    add_bullet(doc, "Gate Rejection with Rework \u2014 Human feedback triggers automatic fix cycles")
    add_bullet(doc, "Director CLI Terminal \u2014 Issue executive directives to agents from dashboard")
    add_bullet(doc, "Approval Management Page \u2014 View and manage all pending human gates")

    # --- Section 5 ---
    add_h1(doc, "5. The Development Workflow")
    add_para(doc,
        "AgentSquad follows a structured 8-phase development workflow, mirroring how a real "
        "software team operates:")

    phases = [
        ("Phase 1: Initialization", "PM spawns core agents. Reads project description. "
         "Seeds Researcher with research topics."),
        ("Phase 2: Research", "Researcher conducts multi-turn technical research. "
         "Produces Research.md with technology evaluation and feasibility analysis."),
        ("Phase 3: Architecture", "PM writes PMSpec.md (business spec with user stories and "
         "acceptance criteria). Architect designs system via 5-turn AI conversation. Produces "
         "Architecture.md. SE reviews architecture for feasibility."),
        ("Phase 4: Engineering Planning", "SE decomposes Architecture into tasks with dependencies "
         "and complexity ratings. PM proposes team composition (core agents + SME specialists). "
         "Human gate approves team. EngineeringPlan.md and TeamComposition.md produced."),
        ("Phase 5: Parallel Development", "SE assigns tasks to engineers based on complexity. "
         "Engineers create PRs with implementation (local build verification). SE + Architect "
         "review PRs. SME agents provide specialist input on-demand. Rework cycles for feedback."),
        ("Phase 6: Testing", "Test Engineer scans approved PRs. Generates three-tier test strategy "
         "(unit \u2192 integration \u2192 UI/E2E). Runs tests locally with Playwright. Classifies "
         "failures as test bugs vs source bugs. Routes rework to appropriate agent."),
        ("Phase 7: Review", "PM conducts final review for business alignment. All PRs pass "
         "multi-agent review pipeline. Human gate for final PR approval."),
        ("Phase 8: Finalization", "SE creates integration PR (final assembly). All PRs merged. "
         "Project marked complete. Workspaces cleaned up."),
    ]
    for phase_name, description in phases:
        add_bold_para(doc, f"{phase_name}: ", description)

    # --- Section 6 ---
    add_h1(doc, "6. Human-Agent Collaboration Model")
    add_para(doc,
        "AgentSquad's most important design principle is that it's built for human-agent "
        "collaboration, not full replacement. The system supports three collaboration modes:")

    add_h2(doc, "Solo Director Mode")
    add_para(doc,
        "One developer manages the entire agent team. They review specs, approve architectures, "
        "guide implementation decisions, and manage quality gates. The developer provides "
        "approximately 15% of the effort (direction, decisions, judgment) while AI handles 85% "
        "(research, writing, coding, testing, review mechanics). This mode is ideal for:")
    add_bullet(doc, "Solo developers who want to operate at team scale")
    add_bullet(doc, "Rapid prototyping where one person needs to ship fast")
    add_bullet(doc, "Proof-of-concept projects with limited staffing")

    add_h2(doc, "Role Pairing Mode")
    add_para(doc,
        "Each human team member pairs with their agent counterpart. The human PM directs the "
        "PM agent on requirements priorities. The human architect reviews the Architect agent's "
        "designs. Human engineers guide their engineer agents on implementation approach. "
        "Each person amplifies their output 5-10x through their AI partner. This mode is ideal for:")
    add_bullet(doc, "Established teams wanting to multiply throughput")
    add_bullet(doc, "Complex projects requiring human domain expertise")
    add_bullet(doc, "Organizations transitioning to AI-augmented development")

    add_h2(doc, "Hybrid Team Mode")
    add_para(doc,
        "A mix of both \u2014 some roles filled by human-agent pairs, others fully automated. "
        "For example, a human architect + human PM pair with their agents for critical decisions, "
        "while engineering and testing run fully autonomous. A small 3-person team could operate "
        "with the throughput of a 15-person team. This mode is ideal for:")
    add_bullet(doc, "Resource-constrained teams with specific expertise gaps")
    add_bullet(doc, "Projects where certain roles need more human judgment than others")
    add_bullet(doc, "Gradual adoption \u2014 automate low-risk roles first, expand over time")

    add_h2(doc, "The Director's Toolkit")
    add_para(doc, "The human director has several tools for guiding the team:")
    add_bullet(doc, "Human Gates \u2014 Approve/reject at any workflow checkpoint with feedback")
    add_bullet(doc, "Director CLI \u2014 Issue natural-language directives to specific agents")
    add_bullet(doc, "Agent Chat \u2014 Have conversations with individual agents for guidance")
    add_bullet(doc, "Configuration \u2014 Adjust model tiers, gate presets, scaling limits at runtime")
    add_bullet(doc, "Prompt Templates \u2014 Edit agent behavior through the dashboard prompt editor")
    add_bullet(doc, "Team Composition \u2014 Review and modify the PM's proposed team before spawning")

    # --- Section 7 ---
    add_h1(doc, "7. Dashboard & Monitoring")
    add_para(doc,
        "The Blazor Server dashboard provides 15 pages of real-time visibility into the agent "
        "team. It runs embedded in the Runner (port 5050) or as a standalone process (port 5051).")

    add_table(doc,
        ["Page", "Description"],
        [
            ["Agent Overview", "Grid of all agents with status badges, model selectors, chat, "
             "error tracking, deadlock alerts"],
            ["Project Timeline", "Visual workflow timeline with PM/Engineering views, phase grouping, "
             "PR/Issue type indicators"],
            ["Metrics", "System health, utilization ring chart, status breakdown, longest-running tasks"],
            ["Health Monitor", "Real-time health checks, stuck agent detection, system diagnostics"],
            ["Pull Requests", "GitHub PR browser with state filters, labels, branch info"],
            ["Issues", "GitHub issue browser with label/assignee filters and sorting"],
            ["Engineering Plan", "Interactive Cytoscape.js dependency graph of engineering tasks"],
            ["Team View", "Visual office-metaphor layout with agent desks and connection lines"],
            ["Director CLI", "Terminal interface for issuing executive directives to agents"],
            ["Approvals", "Human gate approval management with filter buttons"],
            ["Configuration", "Settings editor, gate presets, SME management, GitHub cleanup"],
            ["Agent Detail", "Deep dive into a single agent with pause/resume/terminate controls"],
            ["Agent Reasoning", "View agent decision-making chains and AI conversation history"],
            ["GitHub Feed", "Live feed of GitHub activity across the project"],
            ["Repository", "Browse repository file tree and content"],
        ])

    # --- Section 8 ---
    add_h1(doc, "8. AI Provider Strategy")
    add_para(doc,
        "AgentSquad uses a four-tier model strategy that assigns model quality based on "
        "the criticality of each agent's decisions:")

    add_table(doc,
        ["Tier", "Default Provider", "Used By", "Rationale"],
        [
            ["Premium", "Claude Opus 4.6", "PM, Architect, SE",
             "Quality-critical decisions: specs, architecture, code review"],
            ["Standard", "Claude Sonnet 4.6", "Researcher, SE, TE",
             "Best cost/quality for research and code generation"],
            ["Budget", "GPT-5.2", "Software Engineers",
             "Cost-effective for low-complexity tasks"],
            ["Local", "Ollama (qwen2.5-coder:14b)", "Software Engineers (alt)",
             "Free, offline, good for simple tasks"],
        ])

    add_para(doc,
        "All tiers route through the GitHub Copilot CLI by default \u2014 no API keys needed. "
        "If the Copilot CLI is unavailable, the system automatically falls back to direct API "
        "providers. MCP tool servers are passed through to the Copilot CLI process for tool-augmented "
        "generation.")

    # --- Section 9 ---
    add_h1(doc, "9. Quality & Safety Engineering")

    add_h2(doc, "Code Quality")
    add_bullet(doc, "Local Build Verification \u2014 No code reaches GitHub until it compiles and passes tests")
    add_bullet(doc, "Multi-Agent Review \u2014 Every PR reviewed by 2+ agents for different quality dimensions")
    add_bullet(doc, "Incremental Modification Rules \u2014 Agents preserve existing code when adding features")
    add_bullet(doc, "Duplicate Type Detection \u2014 SE review catches duplicate class/type definitions")
    add_bullet(doc, "Namespace Consistency \u2014 Engineers use namespaces matching existing code structure")

    add_h2(doc, "Operational Safety")
    add_bullet(doc, "Deadlock Detection \u2014 DFS wait-for graph detects circular agent dependencies")
    add_bullet(doc, "Stuck Agent Detection \u2014 Health monitor flags agents idle beyond threshold")
    add_bullet(doc, "Rate Limit Protection \u2014 Proactive throttling prevents GitHub API exhaustion")
    add_bullet(doc, "Graceful Shutdown \u2014 State persisted to SQLite, clean agent lifecycle")
    add_bullet(doc, "MCP Security Policy \u2014 Blocks dangerous server commands, enforces HTTPS, rejects private network URLs")
    add_bullet(doc, "SME Spawn Limits \u2014 Hard caps prevent runaway agent creation")

    add_h2(doc, "Crash Recovery")
    add_bullet(doc, "SQLite State Persistence \u2014 Agent state, gate approvals, memory survive restarts")
    add_bullet(doc, "GitHub as Source of Truth \u2014 PRs/Issues/comments recover lost in-memory state")
    add_bullet(doc, "CLI Session Persistence \u2014 Copilot CLI session IDs resume AI conversations after restart")
    add_bullet(doc, "Boot-Time Filtering \u2014 Dashboard shows only agents from current run, not stale historical data")

    # --- Section 10 ---
    add_h1(doc, "10. How to Use AgentSquad")

    add_h2(doc, "Prerequisites")
    add_bullet(doc, ".NET 8 SDK or later")
    add_bullet(doc, "GitHub Personal Access Token with repo scope")
    add_bullet(doc, "GitHub Copilot CLI v1.0.18+ (default AI provider \u2014 no API keys needed)")
    add_bullet(doc, "Or: at least one AI provider API key (Anthropic, OpenAI, or local Ollama)")

    add_h2(doc, "Getting Started")
    add_bold_para(doc, "Step 1: ", "Clone the repository and run dotnet build")
    add_bold_para(doc, "Step 2: ", "Configure appsettings.json with your project name, description, "
                  "and GitHub repo. Store your GitHub PAT via dotnet user-secrets")
    add_bold_para(doc, "Step 3: ", "Run the Runner (dotnet run from src/AgentSquad.Runner). Agents "
                  "begin the development lifecycle automatically")
    add_bold_para(doc, "Step 4: ", "Monitor via the dashboard at http://localhost:5050 (embedded) "
                  "or http://localhost:5051 (standalone)")
    add_bold_para(doc, "Step 5: ", "Interact as Director: approve gates, reject with feedback, "
                  "issue directives, adjust configuration as needed")

    add_h2(doc, "Configuration")
    add_para(doc,
        "All configuration lives under the AgentSquad section in appsettings.json. Key sections: "
        "Project (GitHub repo/PAT), Models (provider/tier definitions), Agents (per-role tier "
        "assignments, MCP servers, knowledge links), Limits (scaling caps, timeouts), Gates "
        "(human gate presets), Dashboard (port, SignalR toggle), SmeAgents (templates, limits).")

    # --- Section 11 ---
    add_h1(doc, "11. Benefits & Value Proposition")

    add_h2(doc, "For Individual Developers")
    add_bullet(doc, "Operate at team scale \u2014 one person directs 7+ AI agents")
    add_bullet(doc, "Focus on decisions, not execution \u2014 AI handles the volume work")
    add_bullet(doc, "Systematic quality \u2014 multi-agent review catches what one person misses")
    add_bullet(doc, "No context switching \u2014 agents maintain state across work sessions")

    add_h2(doc, "For Engineering Teams")
    add_bullet(doc, "3-5x throughput multiplication for existing teams")
    add_bullet(doc, "Consistent quality through automated review pipelines")
    add_bullet(doc, "Knowledge preservation \u2014 agent memory captures decisions and lessons")
    add_bullet(doc, "Reduced onboarding time \u2014 new team members pair with trained agents")
    add_bullet(doc, "Flexible adoption \u2014 automate one role at a time, expand gradually")

    add_h2(doc, "For Engineering Leaders")
    add_bullet(doc, "Address staffing constraints without hiring")
    add_bullet(doc, "Reduce backlog with AI-augmented throughput")
    add_bullet(doc, "Full audit trail \u2014 every decision in GitHub, every reasoning chain logged")
    add_bullet(doc, "Configurable risk management via graduated autonomy")
    add_bullet(doc, "Platform investment \u2014 reusable across projects and teams")

    # --- Section 12 ---
    add_h1(doc, "12. Competitive Differentiation")
    add_para(doc,
        "AgentSquad occupies a unique position in the AI-assisted development landscape:")

    add_h2(doc, "vs. Single-Agent Coding Tools (Copilot Workspace, Cursor, Cline)")
    add_para(doc,
        "These tools augment individual developers. AgentSquad augments entire teams. Single-agent "
        "tools lack role specialization, peer review, workflow coordination, and test automation. "
        "AgentSquad produces reviewed, tested code through real GitHub workflows \u2014 not just "
        "generated code in an editor.")

    add_h2(doc, "vs. Gas Town (Multi-Agent Orchestrator)")
    add_para(doc,
        "Gas Town (gastownhall/gastown) is the closest comparable system \u2014 a Go-based "
        "multi-agent workspace manager supporting 20-50+ generic worker agents across 10+ AI "
        "runtimes. Key differences:")
    add_table(doc,
        ["Dimension", "AgentSquad", "Gas Town"],
        [
            ["Philosophy", "Opinionated SDLC pipeline with role specialization",
             "Flexible workspace manager with generic workers"],
            ["Agents", "7 fixed roles mirroring real team", "Generic workers assigned tasks"],
            ["Workflow", "Phase-gated linear progression", "Convoy-based flexible orchestration"],
            ["AI Runtime", "Copilot CLI (unified)", "10+ runtimes (Claude Code, Codex, Cursor, etc.)"],
            ["Scale", "7-10 agents (focused)", "20-50+ agents (swarm-scale)"],
            ["Dashboard", "15-page Blazor (Grafana-style)", "htmx web + TUI terminal"],
            ["Language", "C# / .NET 8", "Go 1.25"],
        ])
    add_para(doc,
        "AgentSquad excels at structured SDLC projects where role specialization and systematic "
        "review matter. Gas Town excels at scale and runtime flexibility. The two approaches are "
        "complementary \u2014 AgentSquad could use Gas Town as a coordination layer, and Gas Town "
        "could adopt AgentSquad's SDLC pipeline as a workflow formula.")

    add_h2(doc, "vs. Enterprise AI Platforms (Devin, Factory AI, Sweep)")
    add_para(doc,
        "Enterprise platforms are proprietary, cloud-hosted, and charge per-task. AgentSquad is "
        "self-hosted, open-architecture, and uses existing Copilot licenses. It gives teams full "
        "control over agent behavior via externalized prompts, full audit trails via GitHub, "
        "and no vendor lock-in.")

    # --- Section 13 ---
    add_h1(doc, "13. Vision & Roadmap")

    add_h2(doc, "The 85/15 Development Model")
    add_para(doc,
        "AgentSquad's vision is a development model where AI handles 85% of execution and humans "
        "provide 15% direction. This is not about replacing developers \u2014 it's about evolving "
        "the developer role from maker to director. The human's judgment, creativity, and domain "
        "expertise become more valuable, not less, when amplified by an AI team.")

    add_h2(doc, "Near-Term Enhancements")
    add_bullet(doc, "Multi-channel notifications (Teams Adaptive Cards, email, browser push)")
    add_bullet(doc, "Interactive agent chat via dashboard for real-time guidance")
    add_bullet(doc, "SE fleet-style parallelism for faster task execution")
    add_bullet(doc, "Cross-project agent memory sharing")
    add_bullet(doc, "Plugin architecture for custom workflow stages")

    add_h2(doc, "Long-Term Vision")
    add_bullet(doc, "Organization-wide agent team management platform")
    add_bullet(doc, "Agent-to-agent learning \u2014 insights from one project improve future projects")
    add_bullet(doc, "Self-improving prompts \u2014 agents measure their own output quality and refine behavior")
    add_bullet(doc, "Cross-repository coordination \u2014 agent teams working across microservices")
    add_bullet(doc, "Azure-hosted AgentSquad-as-a-Service for enterprise deployment")

    add_para(doc,
        "The AI-augmented development model represents the next evolution of software engineering. "
        "AgentSquad is the working prototype that proves it's possible today.")

    doc.add_paragraph()
    add_para(doc, "For questions or discussion, contact Ben Humphrey (behumphr@microsoft.com)")

    path = os.path.join(DOCS_DIR, "AgentSquad-Complete-Solution-Guide.docx")
    doc.save(path)
    print(f"Detailed Solution Guide saved to {path}")


if __name__ == "__main__":
    generate_executive_summary()
    generate_detailed_summary()
    print("\nBoth documents generated successfully!")
